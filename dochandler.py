from docx import Document
import time
import json
import re

import osutils
import provider

SETTINGS_ENCODING = "utf-8"
DELAY_BETWEEN_CALLS = 0.1 # in seconds

class DocHandler:
    """
    Receives a URI of word doc and determines if it can translate it.
    Currently supports to translate 1 doc at a time.
    """
    

    def __init__(self, suffix:str="-ES", src:str='Koreano', dst:str='Español') -> None:
        "Make sure its a Word doc. Throws an error if not."
        
        self._doing_translation = False
        self._doc = None
        self._input_docname = ""
        self._output_docname = ""
        self._outputpath = ""
        self.uri = ""
        self.provider = provider.Translator()

        with open("docsettings.json", encoding=SETTINGS_ENCODING) as f:
            self.docsettings = json.load(f)
        
        self.supported_providers = self.docsettings["supported_providers"]
        self.selected_provider = self.docsettings["selected_provider"]
        self.supported_languages = list(self.docsettings["lang_with_codes"].keys())
        self.langdict = self.docsettings["lang_with_codes"]
        self.from_lang = self.docsettings["from_lang"]
        self.to_lang = self.docsettings["to_lang"]
        self.suffix = self.docsettings["suffix"]
    
    def save_docsettings(self,):
        # acquire the current values
        self.docsettings["selected_provider"] = self.selected_provider
        self.docsettings["from_lang"] = self.from_lang
        self.docsettings["to_lang"] = self.to_lang
        if self.suffix == "": # protection in case of empty suffix (to not replace original)
            self.suffix = "-suffix-necesary"
        self.docsettings["suffix"] = self.suffix

        with open("docsettings.json", "w", encoding=SETTINGS_ENCODING) as f:
            json.dump(self.docsettings, f, indent=4)
        
        # new settings changed, so update outputname if doc is already loaded
        if self.uri:
            self.set_document(self.uri)

    def set_document(self, uri:str):
        self.uri = uri
        self._path = osutils.parse_uri2path(uri)
        folder, self._input_docname = osutils.get_folder_n_child_from_uri(uri)
        self._doc = Document(self._path)
        self._outputpath = osutils.get_filepath_wo_ext(self._path)+self.suffix+".docx"
        self._output_docname = osutils.get_filepath_wo_ext(self._input_docname)+self.suffix+".docx"
    
    # feedback the output name of the file translated
    # def set_lbfeedback(self, lbfeedback):
    #     self.lbfeedback = lbfeedback

    def can_translate(self,):
        if not self._doc:
            return True
        elif not self._doing_translation:
            return True
        else:
            return False

    def start_translation(self, glib_idle_add, update_ui_func, success_cb, needs_check_cb):
        """
        This functions is called from a thread, so preferably not return any value.
        Starts translating while updating UI with the function pointers.
        `glib_idel_add uses` uses `update_ui_func` sending a parameter for it.
        e.g.:
        ```python
        def update_progess(i):
            progress.pulse()
            progress.set_text(str(i))
            return False

        def example_target():
            for i in range(50):
                GLib.idle_add(update_progess, i)
                time.sleep(0.2)
        ```
        """
        self._doing_translation = True
        status_ok = True
        no_paragraphs = len(self._doc.paragraphs)
        for i in range(no_paragraphs):
            line = self._doc.paragraphs[i].text.strip()
            # if line starts and ends with " and contains spaces + dots
            # https://www.programiz.com/python-programming/regex
            if re.match('^"\s*\.+\s*"$', line):
                continue # skip to next iteration
            if "CAP" in line:
                continue
            elif "***" == line:
                continue
            else:
                translated_text:str = self.from_service(line)
                if translated_text == "429":
                    print("Character limit reached!")
                    status_ok = False
                    break
                elif translated_text == "401":
                    print("API keys are incorrect")
                    status_ok = False
                    break
                elif len(translated_text) > 3:
                    self._doc.paragraphs[i].text = translated_text
                    message = f"Párrafo {i+1} de {no_paragraphs}"
                    glib_idle_add(update_ui_func, message)
                    time.sleep(DELAY_BETWEEN_CALLS)
        
        self._doc.save(self._outputpath)
        success_cb() if status_ok else needs_check_cb()
        self._doing_translation = False

    def do_start_translation(self, glib_idle_add, update_ui_func, success_cb, needs_check_cb):
        "This function is a simulation of `start_translation`"
        self._doing_translation = True
        status_ok = True
        no_paragraphs = len(self._doc.paragraphs)
        for i in range(no_paragraphs):
            line = self._doc.paragraphs[i].text.strip()
            if "CAP" in line:
                continue
            elif "***" == line:
                continue
            else:
                # conditional to check either case (completed or incompleted translation)
                if i < no_paragraphs-50:
                    message = f"Párrafo {i+1} de {no_paragraphs}"
                    glib_idle_add(update_ui_func, message)
                    time.sleep(0.3)
                else:
                    status_ok = False
                    break

        self._doc.save(self._outputpath)
        success_cb() if status_ok else needs_check_cb()
        self._doing_translation = False

    def from_service(self, text:str,) -> bool:
        src = self.langdict[self.from_lang]
        dst = self.langdict[self.to_lang]
        if self.selected_provider == "Papago":
            translated_text = self.provider.use_papago(text, src, dst)
        elif self.selected_provider == "Kakao":
            translated_text = self.provider.use_kakao(text, src, dst)
        elif self.selected_provider == "Google":
            translated_text = self.provider.use_google(text, src, dst)
        return translated_text
    
    def get_filename(self,)->str:
        return self._input_docname

    def get_output_filename(self,)->str:
        return self._output_docname



def is_a_word_doc(uri:str):
    if osutils.is_file_or_folder_uri(uri): # its a file
        folder, docname = osutils.get_folder_n_child_from_uri(uri)
        if docname.endswith(".docx") or docname.endswith(".doc"):
            # print("URI:", uri, "is valid")
            return True
        else:
            # it may be a PDF which we can be a feature in future versions
            print(docname, "is not a Word document")
            return False

    else: # its a folder
        # another possible feature is read folders
        # parent, child = osutils.get_folder_n_child_from_uri(uri)
        # print("Parent folder:", parent)
        # print("Child folder:", child)
        return False
